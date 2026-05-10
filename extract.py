"""
Core metadata extraction using Google Gemini (multimodal LLM).

Supports:
  - .docx / .pdf.docx  →  text extracted via python-docx, sent to Gemini
  - .png / .jpg        →  image sent directly to Gemini (vision)
"""

import os
import re
import json
import time
from pathlib import Path

from dotenv import load_dotenv
from google import genai
from docx import Document
from PIL import Image

# ---------------------------------------------------------------------------
# Configure Gemini
# ---------------------------------------------------------------------------
load_dotenv()
API_KEY = os.environ["GOOGLE_API_KEY"]
client = genai.Client(api_key=API_KEY)
MODEL = "gemini-flash-latest"

# ---------------------------------------------------------------------------
# Prompt
# ---------------------------------------------------------------------------
PROMPT = """You are an expert at reading rental/lease agreement documents.

Extract these 6 fields and return ONLY a valid JSON object — no markdown, no explanation:

{
  "Aggrement Value": <monthly rent as an integer, no currency symbol>,
  "Aggrement Start Date": <date as "DD.MM.YYYY">,
  "Aggrement End Date": <date as "DD.MM.YYYY">,
  "Renewal Notice (Days)": <integer number of days, or null if not mentioned>,
  "Party One": <first party name — usually the landlord / owner / lessor>,
  "Party Two": <second party name — usually the tenant / lessee>
}

Strict rules:
- "Aggrement Value": plain integer only (e.g. 6500, never "Rs. 6,500").
- "Aggrement Start Date" / "Aggrement End Date": DD.MM.YYYY format. If a date like "31st April" appears, use the last valid day of that month (30.04.YYYY).
- "Renewal Notice (Days)": always convert to an integer number of days. Examples: "2 months" = 60, "3 months" = 90, "1 month" = 30, "15 days" = 15. If not mentioned at all, return null.
- "Party One" / "Party Two": return ONLY the bare name — NO titles, honorifics, or prefixes of any kind (no Mr., Mrs., Ms., Dr., Prof., Sri., Smt., Shri., M/s, etc.). If a party has multiple people, join with " & ".
- If a field genuinely cannot be found, use null.
- If the document contains multiple agreements, extract from the PRIMARY (main) agreement only.
- Return NOTHING except the JSON object.
"""

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _read_docx(path: Path) -> str:
    """Extract all text from a .docx file including table cells."""
    doc = Document(path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _parse_response(raw: str) -> dict:
    """Strip any markdown fencing and parse JSON from LLM response."""
    text = raw.strip()
    fenced = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    if fenced:
        text = fenced.group(1).strip()
    return json.loads(text)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_from_file(file_path) -> dict:
    """
    Extract rental agreement metadata from a .docx or image file.

    Returns a dict with the 6 target fields.
    """
    file_path = Path(file_path)
    name_lower = file_path.name.lower()
    suffix = file_path.suffix.lower()

    if suffix == ".docx" or name_lower.endswith(".pdf.docx"):
        text = _read_docx(file_path)
        contents = PROMPT + f"\n\nDocument text:\n{text}"

    elif suffix in {".png", ".jpg", ".jpeg"}:
        image = Image.open(file_path)
        contents = [PROMPT, image]

    else:
        raise ValueError(f"Unsupported file format: {suffix}")

    for attempt in range(3):
        try:
            response = client.models.generate_content(model=MODEL, contents=contents)
            return _parse_response(response.text)
        except Exception as e:
            if attempt < 2 and ("503" in str(e) or "UNAVAILABLE" in str(e)):
                time.sleep(20)
                continue
            raise


# ---------------------------------------------------------------------------
# CLI usage: python extract.py <file_path>
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python extract.py <path_to_document>")
        sys.exit(1)

    result = extract_from_file(sys.argv[1])
    print(json.dumps(result, indent=2))
